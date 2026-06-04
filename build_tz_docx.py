# -*- coding: utf-8 -*-
"""«Dasturxon» texnik topshirig'ini professional Word (.docx) hujjat sifatida yaratadi."""
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x12, 0x20, 0x36)
NAVY2  = RGBColor(0x1B, 0x2F, 0x4D)
ACCENT = RGBColor(0x2F, 0x5F, 0xA6)
MUTED  = RGBColor(0x5C, 0x66, 0x75)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
FAINT  = RGBColor(0x8A, 0x93, 0xA3)
NAVY_HEX, ACCENT_HEX, SOFT_HEX = "1B2F4D", "2F5FA6", "EAF0F9"

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0x1A, 0x22, 0x2E)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz, col in (('Heading 1', 15, NAVY2), ('Heading 2', 12.5, ACCENT), ('Heading 3', 11, NAVY2)):
    st = doc.styles[lvl]
    st.font.name = 'Calibri'
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True

# ---- A4 + margins ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin = sec.bottom_margin = Mm(20)
sec.left_margin = sec.right_margin = Mm(22)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    for i, part in enumerate(str(text).split('\n')):
        run = (p if i == 0 else p.add_run('\n')) and p.add_run(part)
        run.bold = bold; run.font.size = Pt(size)
        if color: run.font.color.rgb = color


def set_widths(table, widths_mm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_mm):
            row.cells[idx].width = Mm(w)


def make_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_text(c, h, bold=True, color=WHITE, size=9.5)
        shade(c, NAVY_HEX)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cell_text(cells[i], val, size=9.5)
    set_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_field(paragraph, instr_text, placeholder=""):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); run._r.append(b)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr_text; run._r.append(it)
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate'); run._r.append(sep)
    if placeholder:
        t = OxmlElement('w:t'); t.text = placeholder; run._r.append(t)
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); run._r.append(e)


def bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p


# =================== COVER PAGE ===================
para("TEXNIK TOPSHIRIQ  ·  ТЕХНИЧЕСКОЕ ЗАДАНИЕ (ТЗ)", color=ACCENT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    doc.add_paragraph()
title = para("«Dasturxon»", align=WD_ALIGN_PARAGRAPH.CENTER)
tr = title.runs[0]; tr.font.size = Pt(40); tr.bold = True; tr.font.color.rgb = NAVY
sub = para("Restoran operatsiyalarini avtomatlashtirish tizimi", align=WD_ALIGN_PARAGRAPH.CENTER)
sub.runs[0].font.size = Pt(14); sub.runs[0].font.color.rgb = MUTED
sub2 = para("Backend · Web boshqaruv paneli · Mobil ilova · QR self-service", align=WD_ALIGN_PARAGRAPH.CENTER)
sub2.runs[0].font.size = Pt(11); sub2.runs[0].font.color.rgb = ACCENT
for _ in range(4):
    doc.add_paragraph()

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ("Hujjat raqami", "DSX-TZ-2026-01"),
    ("Versiya", "1.0"),
    ("Sana", "2026-yil 26-may"),
    ("Maxfiylik darajasi", "Ichki foydalanish uchun"),
    ("Buyurtmachi", "Xolmatov Hasanboy G'ayratjon o'g'li"),
    ("Ishlab chiquvchi", "andijanrestoran-byte"),
    ("Hujjat holati", "Tasdiqlash uchun"),
]
for k, v in meta_rows:
    cells = meta.add_row().cells
    cell_text(cells[0], k, bold=True, color=MUTED, size=10)
    cell_text(cells[1], v, size=10)
set_widths(meta, [55, 95])
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# =================== DOCUMENT CONTROL ===================
h1("Hujjat nazorati")
make_table(["Versiya", "Sana", "Tavsif", "Muallif"],
           [["1.0", "2026-05-26", "Hujjatning birlamchi (boshlang'ich) versiyasi.", "Loyiha jamoasi"]],
           [22, 30, 78, 36])

# =================== TOC ===================
h1("Mundarija")
para("Mundarijani yangilash: ustiga bosib, F9 (yoki o'ng tugma → Update Field).", italic=True, color=FAINT, size=9)
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u', "«Update Field» bosilganda mundarija shakllanadi.")
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# =================== 1 ===================
h1("1. Umumiy ma'lumotlar")
para("Ushbu texnik topshiriq «Dasturxon» dasturiy tizimiga qo'yiladigan talablarni belgilaydi. "
     "«Dasturxon» — restoran (oshxona/kafe) kundalik operatsiyalarini avtomatlashtirishga mo'ljallangan "
     "yaxlit tizim bo'lib, to'rt turdagi foydalanuvchini — direktor, kassir, ofitsant va mijoz — yagona "
     "backend va ma'lumotlar bazasi atrofida birlashtiradi.")
para("Tizim quyidagi o'zaro bog'liq komponentlardan tashkil topadi:")
make_table(["Komponent", "Vazifasi", "Texnologiya"], [
    ["Backend (markaziy server)", "Biznes-mantiq, ma'lumotlar bazasi, REST xizmatlari", "Django 6 + DRF"],
    ["Web boshqaruv paneli", "Kassir va direktorning asosiy ish joyi", "Django shablonlari"],
    ["Mobil ilova", "Ofitsant/direktor uchun buyurtma olish", "Flutter"],
    ["Mijoz veb sahifasi", "QR orqali o'z-o'ziga xizmat (autentifikatsiyasiz)", "Django (public)"],
], [42, 78, 46])

# =================== 2 ===================
h1("2. Maqsad va vazifalar")
para("Asosiy maqsad. Restoranda buyurtma qabul qilish, oshxonaga uzatish, to'lovni rasmiylashtirish va "
     "moliyaviy hisobotlarni shakllantirish jarayonlarini to'liq raqamlashtirish hamda inson omilidan "
     "kelib chiqadigan xatolarni kamaytirish.")
para("Vazifalar:")
bullets([
    "ofitsant va mijozdan buyurtmalarni real vaqtda qabul qilish;",
    "buyurtmalarni avtomatik ravishda oshxonaga (printerga) yo'naltirish;",
    "bitta stolda bir nechta mustaqil hisob (multi-bill) yuritish;",
    "kunlik portsiya (taom qoldig'i) hisobini yuritish;",
    "naqd va karta to'lovlarini qayd etish hamda hisobni yopish;",
    "direktorga kunlik/haftalik/oylik tushum va xodimlar unumdorligi bo'yicha hisobot berish;",
    "QR orqali mijozga o'z-o'ziga xizmat ko'rsatish imkonini berish.",
])

# =================== 3 ===================
h1("3. Atamalar va qisqartmalar")
make_table(["Atama", "Izoh"], [
    ["Smena (Shift)", "Bir kunlik ish davri. Smena ochilmaguncha buyurtma qabul qilish va portsiya kiritish bloklanadi."],
    ["Shot (bill_number)", "Bitta stoldagi mustaqil hisob raqami (1–10). Bir stolga 10 tagacha alohida hisob."],
    ["Portsiya", "Taomning shu kunga belgilangan miqdori; iste'mol qilingani avtomatik ayriladi."],
    ["is_rejectable", "Taomni rad etish mumkinligi belgisi. Qiymati False bo'lsa, taom rad etilmaydi."],
    ["Oshxona konsoli", "Oshxona printeriga ulangan kompyuterda ochiq turuvchi, cheklarni avtomatik chop etuvchi sahifa."],
    ["QR self-service", "Mijozning telefon orqali QR-kodni skanerlab, menyu ko'rishi va buyurtma berishi."],
    ["RBAC", "Rolga asoslangan kirish nazorati (Role-Based Access Control)."],
    ["DRF / JWT", "Django REST Framework / JSON Web Token."],
], [42, 124])

# =================== 4 ===================
h1("4. Loyiha doirasi")
h2("4.1. Tizimga kiritiladi")
bullets([
    "autentifikatsiya (web — sessiya, mobil/tashqi klient — JWT) va rolga asoslangan kirish nazorati;",
    "menyu va kategoriyalarni boshqarish (CRUD), kunlik portsiya hisobi;",
    "stollar, QR-kodlar va multi-bill hisoblar;",
    "buyurtma qabul qilish (ofitsant mobil + mijoz QR) va avtomatik ACCEPTED holati;",
    "oshxona konsoli orqali avtomatik chek chop etish;",
    "smena boshqaruvi va gating mantiq;",
    "to'lov (naqd/karta) va hisobni yopish;",
    "moliyaviy hisobotlar va boshqaruv paneli.",
])
h2("4.2. Tizimga kiritilmaydi (joriy versiyada)")
bullets([
    "onlayn to'lov shlyuzlari (Click, Payme va h.k.) bilan integratsiya;",
    "ombor/ta'minot (inventarizatsiya) to'liq moduli;",
    "sodiqlik (loyalty) dasturi va chegirma kartalari;",
    "bir nechta filial (multi-branch) boshqaruvi.",
])

# =================== 5 ===================
h1("5. Foydalanuvchi rollari")
make_table(["Rol", "Kod", "Tavsif"], [
    ["Direktor", "director", "To'liq nazorat: xodimlar (ofitsant) CRUD, moliyaviy hisobotlar, naqd/karta savdolari. Web panelga teng huquqda kira oladi."],
    ["Kassir", "cashier", "Web paneldagi asosiy ish o'rni: smena, menyu va portsiya, buyurtmalar, stollar, hisobni yopish va to'lov, oshxona konsoli."],
    ["Ofitsant", "waiter", "Mobil ilova orqali stolga birikadi va buyurtma oladi (stol + shot). Buyurtmasi avtomatik ACCEPTED bo'ladi."],
    ["Mijoz", "public", "Autentifikatsiyasiz. QR orqali menyuni ko'radi, savat to'ldiradi, buyurtma beradi va holatini kuzatadi."],
], [26, 24, 116])

# =================== 6 ===================
h1("6. Funksional talablar")
para("Har bir talab noyob identifikator bilan belgilangan.", italic=True, color=MUTED, size=9.5)

fr_groups = [
    ("6.1. Autentifikatsiya va avtorizatsiya", True, [
        ("FR-AUTH-01", "Web foydalanuvchi sessiya orqali, mobil/tashqi klient esa JWT (Bearer token) orqali kirishi shart.", "Yuqori"),
        ("FR-AUTH-02", "Kirish jarayoni login va parolni qabul qilib, foydalanuvchini roli bilan birga aniqlashi kerak.", "Yuqori"),
        ("FR-AUTH-03", "Har bir amal foydalanuvchi roliga qarab cheklanishi (RBAC) shart: director, cashier, waiter.", "Yuqori"),
        ("FR-AUTH-04", "Sessiya/token muddati tugaganda foydalanuvchi qayta kirishga yo'naltiriladi.", "O'rta"),
    ]),
    ("6.2. Direktor moduli", False, [
        ("FR-DIR-01", "Ofitsantlar ustida CRUD (yaratish / tahrirlash / o'chirish)."),
        ("FR-DIR-02", "Kunlik / haftalik / oylik tushum hisobotlari."),
        ("FR-DIR-03", "Ofitsantlar unumdorligi hisoboti (kim qancha sotgan, qaysi taomlar)."),
        ("FR-DIR-04", "Naqd va karta savdolari to'lov yozuvlaridan avtomatik agregatlanadi."),
    ]),
    ("6.3. Kassir moduli (web panel)", False, [
        ("FR-CASH-01", "Boshqaruv paneli — barcha bo'limlarga yagona kirish nuqtasi."),
        ("FR-CASH-02", "Buyurtmalar ro'yxati — faqat faol buyurtmalar karta ko'rinishida (Stol №N, summa, ofitsant, mahsulot soni)."),
        ("FR-CASH-03", "Buyurtma tafsiloti — alohida pozitsiyani rad etish; is_rejectable=False taom rad etilmaydi."),
        ("FR-CASH-04", "Stollar ekrani — har bir stol va faol hisoblar, ~20 soniyada avtomatik yangilanadi."),
        ("FR-CASH-05", "Stol hisobi — shot bo'yicha; hisobni yopish amali to'lov modalini chaqiradi."),
        ("FR-CASH-06", "Stollar QR ekrani — barcha stollar QR rasmlari, chop etish va yuklab olish."),
        ("FR-CASH-07", "Xodimlar ekrani — ofitsant ro'yxati, qo'shish, o'chirish."),
    ]),
    ("6.4. Ofitsant moduli (mobil ilova)", False, [
        ("FR-WAIT-01", "Ofitsant stolga birikadi va buyurtma oladi (stol + shot raqami 1–10)."),
        ("FR-WAIT-02", "Ofitsant buyurtmasi avtomatik Accepted bo'ladi (kassir tasdiqlashi shart emas)."),
        ("FR-WAIT-03", "Menyuda har taom uchun qolgan porsiya ko'rsatiladi; tugaganda tanlash bloklanadi."),
        ("FR-WAIT-04", "«Rad etilmaydi» belgisi tegishli taomlarda ko'rsatiladi."),
    ]),
    ("6.5. Mijoz moduli (QR self-service)", False, [
        ("FR-CLNT-01", "QR orqali mobil menyu: kategoriyalar, narx, portsiya, savat (+/−) va buyurtma berish (ism + izoh)."),
        ("FR-CLNT-02", "Tugagan taom ko'rsatiladi, ammo tanlab bo'lmaydi."),
        ("FR-CLNT-03", "Buyurtma holatini kuzatish (avtomatik yangilanish), har pozitsiya holati va jami summa."),
        ("FR-CLNT-04", "Mijoz rad etiluvchi taomni rad etishi mumkin; portsiya qaytadi; barchasi rad etilsa buyurtma bekor qilinadi."),
    ]),
    ("6.6. Oshxona konsoli", False, [
        ("FR-KIT-01", "Konsol sahifasi printer ulangan kompyuterda ochiq qoladi va davriy (~8 s) tekshiradi."),
        ("FR-KIT-02", "Yangi (qabul qilingan) buyurtma cheki avtomatik chop etiladi."),
        ("FR-KIT-03", "Takror chop etishning oldi olinadi: kitchen_printed bayrog'i va atomik «claim»."),
    ]),
    ("6.7. Smena tizimi", False, [
        ("FR-SHIFT-01", "Kassir smenani boshlaydi (bir vaqtda bitta ochiq smena; yopib qayta ochish mumkin)."),
        ("FR-SHIFT-02", "Smena ochilmaguncha buyurtma qabul qilib bo'lmaydi (web va API darajasida)."),
        ("FR-SHIFT-03", "Smena ochilmaguncha kunlik portsiya kiritib bo'lmaydi."),
        ("FR-SHIFT-04", "Smena ochilgach portsiya jadvali; qolgan = boshlang'ich − iste'mol."),
        ("FR-SHIFT-05", "Kassir/direktor smenani yopadi (closed_at, closed_by). Yopilgach buyurtma qabul qilinmaydi."),
    ]),
    ("6.8. Menyu va portsiya boshqaruvi", False, [
        ("FR-MENU-01", "Taomlar to'liq CRUD: nom, kategoriya, narx, tavsif, faollik, «Rad etish mumkin»."),
        ("FR-MENU-02", "Kategoriyalar CRUD va kategoriya bo'yicha guruhlash."),
        ("FR-MENU-03", "Kunlik portsiya jadvali: har taom uchun boshlang'ich va qolgan miqdor."),
    ]),
    ("6.9. To'lov va hisobni yopish", False, [
        ("FR-PAY-01", "Hisob yopilganda Naqd / Karta tanlash modali chiqadi."),
        ("FR-PAY-02", "Tanlangach har buyurtma uchun to'lov yozuvi yaratiladi (summa va kassir)."),
        ("FR-PAY-03", "Buyurtma COMPLETED bo'ladi va chek chiqadi; chekda to'lov turi ko'rsatiladi."),
    ]),
]
for title_txt, with_pri, rows in fr_groups:
    h2(title_txt)
    if with_pri:
        make_table(["ID", "Talab", "Prioritet"], [[r[0], r[1], r[2]] for r in rows], [26, 116, 24])
    else:
        make_table(["ID", "Talab"], [[r[0], r[1]] for r in rows], [26, 140])

# =================== 7 ===================
h1("7. Ma'lumotlar modeli")
para("Asosiy entitilar (modellar) va ularning kalit maydonlari.", italic=True, color=MUTED, size=9.5)
make_table(["Model", "Asosiy maydonlar"], [
    ["DiningTable", "qr_token, assigned_waiters, current_status"],
    ["Order", "bill_number (1–10), status, order_source (waiter/client), client_name, public_token, total_amount, payable_amount, kitchen_printed"],
    ["OrderItem", "product, quantity, status (pending/accepted/rejected), rejection_reason, line_total"],
    ["Product", "price, is_active, is_rejectable, image, qolgan portsiya hisobi"],
    ["MenuCategory", "name, sort_order"],
    ["ProductDailyStock", "(product, date) noyob, initial_quantity, remaining_quantity"],
    ["Shift", "date, opened_by, opened_at, closed_by, closed_at (ochiqligi closed_at bilan)"],
    ["Payment", "order (1:1), payment_method (cash/card/mixed), amount, cashier"],
], [38, 128])

# =================== 8 ===================
h1("8. Texnik talablar va arxitektura")
make_table(["Qatlam / yo'nalish", "Texnologiya va talab"], [
    ["Backend freymvork", "Python, Django 6 + Django REST Framework"],
    ["API autentifikatsiya", "JWT (djangorestframework-simplejwt)"],
    ["API hujjatlari", "drf-spectacular (OpenAPI)"],
    ["CORS", "django-cors-headers"],
    ["Ma'lumotlar bazasi", "Ishlab chiqishda SQLite, produksiyada PostgreSQL (psycopg, dj-database-url)"],
    ["Web panel", "Django shablonlari, zamonaviy SaaS dizayn"],
    ["Mobil ilova", "Flutter (Dart)"],
    ["QR generatsiya", "qrcode, Pillow"],
    ["Statik / server", "whitenoise, gunicorn"],
    ["Joylashtirish", "Railway yoki muqobil; migrate → collectstatic → gunicorn"],
], [50, 116])

# =================== 9 ===================
h1("9. Nofunksional talablar")
make_table(["ID", "Kategoriya", "Talab"], [
    ["NFR-01", "Xavfsizlik", "RBAC; JWT bilan himoyalangan API; parollar standart usulda xeshlanadi; CSRF/CORS to'g'ri."],
    ["NFR-02", "Yaxlitlik", "To'lov va portsiya kamayishi atomik; oshxona chop etish atomik «claim»."],
    ["NFR-03", "Real vaqtlilik", "Oshxona ~8 s, stollar ~20 s, mijoz holati ~15 s da yangilanadi."],
    ["NFR-04", "Qulaylik", "Web panel responsive va zamonaviy; mobil ilova sodda va tez."],
    ["NFR-05", "Chop etish", "Chek va QR sahifalari printer uchun oq fonli."],
    ["NFR-06", "Ishonchlilik", "O'zgarishlardan keyin smoke testlar; loyiha «check» dan toza o'tadi."],
    ["NFR-07", "Tillilik", "Foydalanuvchi interfeysi o'zbek tilida."],
], [20, 34, 112])

# =================== 10 ===================
h1("10. Interfeys va dizayn talablari")
bullets([
    "Web panel: zamonaviy SaaS tema — Manrope shrifti, neon-lime accent, ikonali sidebar, guruhlangan navigatsiya.",
    "Buyurtmalar: texnik identifikatorlar o'rniga foydalanuvchiga tushunarli «Stol №N» ko'rinishi.",
    "Holat indikatorlari: tugagan taom «Tugadi»; «Rad etilmaydi» — alohida belgi.",
    "Mobil ilova: Flutter; statik tahlildan toza o'tadi; har taomda qolgan portsiya va rad etish belgisi.",
])

# =================== 11 ===================
h1("11. Cheklovlar va taxminlar")
bullets([
    "bir stolga maksimal 10 ta mustaqil hisob (shot);",
    "bir vaqtda bitta ochiq smena; smenani yopib o'sha kun ichida qayta ochish mumkin;",
    "oshxona konsoli ishlashi uchun sahifa tegishli kompyuterda doimo ochiq turishi shart;",
    "ofitsant va mijoz buyurtmalari avtomatik ACCEPTED bo'ladi (qo'lda qabul talab etilmaydi).",
])
note = doc.add_paragraph()
nr = note.add_run("Xavfsizlik eslatmasi. Git remote URL'da GitHub kirish tokeni (PAT) ochiq qolmasligi kerak — "
                  "tokenni yangilab, SSH yoki credential manager orqali saqlash tavsiya etiladi.")
nr.italic = True; nr.font.color.rgb = RGBColor(0x9A, 0x62, 0x12); nr.font.size = Pt(9.5)

# =================== 12 ===================
h1("12. Rivojlantirish rejasi")
bullets([
    "direktor uchun naqd/karta savdolarini alohida web sahifada ko'rsatish;",
    "mijoz sahifasiga «savatni tozalash» va miqdorni aniq ko'rsatish;",
    "real qurilmada to'liq oqimni (mijoz → kassir) sinovdan o'tkazish;",
    "kelajakda onlayn to'lov shlyuzlari bilan integratsiya.",
])

# =================== SIGN-OFF ===================
h1("Tasdiqlash")
para("Ushbu texnik topshiriq tomonlar tomonidan ko'rib chiqilib, kelishilgan holda tasdiqlanadi.",
     italic=True, color=MUTED, size=9.5)
sign = doc.add_table(rows=1, cols=2)
sign.style = 'Table Grid'
sign.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = [("Buyurtmachi", "Xolmatov Hasanboy G'ayratjon o'g'li"),
          ("Ishlab chiquvchi", "andijanrestoran-byte")]
for i, (role, nm) in enumerate(labels):
    c = sign.rows[0].cells[i]
    c.text = ''
    p1 = c.paragraphs[0]; r1 = p1.add_run(role.upper()); r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = FAINT
    p2 = c.add_paragraph(); r2 = p2.add_run(nm); r2.bold = True; r2.font.size = Pt(10.5)
    c.add_paragraph(); c.add_paragraph()
    p3 = c.add_paragraph(); r3 = p3.add_run("Imzo: ______________      Sana: ____________"); r3.font.size = Pt(9.5); r3.font.color.rgb = MUTED
set_widths(sign, [83, 83])

# =================== FOOTER (page numbers) ===================
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.text = ""
left = fp.add_run("«Dasturxon» — Texnik topshiriq · DSX-TZ-2026-01 · v1.0     |     ")
left.font.size = Pt(8); left.font.color.rgb = FAINT
fp.add_run("Sahifa ").font.size = Pt(8)
add_field(fp, "PAGE", "1")
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in fp.runs:
    r.font.size = Pt(8); r.font.color.rgb = FAINT

# ---- update fields (TOC, page) on open ----
settings = doc.settings.element
uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)

out = r"C:\Users\xasan\PycharmProjects\andijan\texnik_zadacha.docx"
doc.save(out)
print("SAQLANDI:", out)
